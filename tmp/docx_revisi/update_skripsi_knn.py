from collections import Counter
from math import sqrt
from pathlib import Path

from docx import Document


SOURCE = Path("tmp/docx_revisi/Seminar_Isi_Skripsi_Steven_Evaluasi_Sistem_Revisi_system.docx")
OUTPUT = Path("tmp/docx_revisi/Seminar_Isi_Skripsi_Steven_Evaluasi_Sistem_Revisi_SYSTEM_KNN.docx")

FEATURE_NAMES = [
    "Pancasila",
    "Matematika",
    "Bahasa Indonesia",
    "Bahasa Inggris",
    "Minat Mapel",
    "Bakat Kemampuan",
]


def parse_decimal(value):
    return float(str(value).strip().replace(".", "").replace(",", "."))


def encode_kelompok(value):
    text = str(value).strip().upper()
    return 1.0 if "KELOMPOK MAPEL 1" in text or text.endswith(" 1") else 0.0


def fmt(value):
    return f"{float(value):.4f}".replace(".", ",")


def fmt_raw(value):
    return f"{float(value):.2f}".replace(".", ",")


def set_paragraph_text(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_cell_text(cell, text):
    paragraph = cell.paragraphs[0]
    set_paragraph_text(paragraph, str(text))
    for extra in cell.paragraphs[1:]:
        set_paragraph_text(extra, "")


def make_formula(test_norm, train_norm):
    parts = []
    for x_val, y_val in zip(test_norm, train_norm):
        parts.append(f"({fmt(x_val)}-{fmt(y_val)})²")
    return "√(" + " + ".join(parts) + ")"


def normalizer(rows):
    mins = [min(row["raw"][idx] for row in rows) for idx in range(6)]
    maxs = [max(row["raw"][idx] for row in rows) for idx in range(6)]
    ranges = [
        (max_value - min_value) if max_value != min_value else 1
        for min_value, max_value in zip(mins, maxs)
    ]

    def norm(values):
        return [
            (float(values[idx]) - mins[idx]) / ranges[idx]
            for idx in range(6)
        ]

    return mins, maxs, ranges, norm


def rounded(values):
    return [round(float(value), 4) for value in values]


doc = Document(SOURCE)

# Data latih alumni.
training_rows = []
for no, row in enumerate(doc.tables[7].rows[1:], start=1):
    cells = [cell.text.strip() for cell in row.cells]
    raw = [parse_decimal(value) for value in cells[1:5]]
    raw.extend([encode_kelompok(cells[5]), encode_kelompok(cells[6])])
    training_rows.append(
        {
            "no": no,
            "name": cells[0],
            "raw": raw,
            "label": cells[7],
        }
    )

mins, maxs, ranges, normalize = normalizer(training_rows)

# Data uji. Override dua nilai kategorikal agar sama dengan hasil yang dipakai sistem.
test_rows = []
for no, row in enumerate(doc.tables[8].rows[1:], start=1):
    cells = [cell.text.strip() for cell in row.cells]
    raw = [parse_decimal(value) for value in cells[1:5]]
    raw.extend([encode_kelompok(cells[5]), encode_kelompok(cells[6])])

    if cells[0].strip().lower() == "ridho":
        raw[4] = 1.0
        raw[5] = 0.0
        set_cell_text(row.cells[5], "Kelompok Mapel 1")
        set_cell_text(row.cells[6], "Kelompok Mapel 2")

    if "FADIL RENALDI SUKMA" in cells[0].upper():
        raw[4] = 1.0
        raw[5] = 0.0
        set_cell_text(row.cells[5], "Kelompok Mapel 1")
        set_cell_text(row.cells[6], "Kelompok Mapel 2")

    test_rows.append(
        {
            "no": no,
            "name": cells[0],
            "raw": raw,
            "label": "Kelompok Mapel 1" if raw[4] == 1 else "Kelompok Mapel 2",
        }
    )

# Update deskripsi normalisasi agar sesuai dengan implementasi sistem.
set_paragraph_text(
    doc.paragraphs[592],
    "Pada tahap ini, data dinormalisasikan menggunakan metode Min-Max Normalization agar seluruh atribut berada pada skala yang seimbang. Nilai minimum dan maksimum diambil dari data latih alumni, kemudian data uji dinormalisasikan menggunakan acuan minimum dan maksimum tersebut. Normalisasi dilakukan pada atribut nilai Pancasila, Matematika, Bahasa Indonesia, dan Bahasa Inggris. Atribut Minat Mapel dan Bakat Kemampuan dikodekan dalam bentuk biner, yaitu Kelompok Mapel 1 = 1 dan Kelompok Mapel 2 = 0.",
)
set_paragraph_text(
    doc.paragraphs[597],
    "Keterangan: Xnorm adalah nilai hasil normalisasi, X adalah nilai asli data, Xmin adalah nilai minimum atribut pada data latih, dan Xmax adalah nilai maksimum atribut pada data latih.",
)

# Tabel IV.5: min dan max data latih.
minmax_table = doc.tables[10]
for idx in range(6):
    set_cell_text(minmax_table.rows[idx + 1].cells[1], fmt_raw(mins[idx]) if idx < 4 else str(int(mins[idx])))
    set_cell_text(minmax_table.rows[idx + 1].cells[2], fmt_raw(maxs[idx]) if idx < 4 else str(int(maxs[idx])))

# Tabel IV.6: normalisasi data uji.
test_norms = {}
test_table = doc.tables[11]
for test in test_rows:
    norm_values = rounded(normalize(test["raw"]))
    test_norms[test["name"]] = norm_values
    row = test_table.rows[test["no"]]
    for idx, value in enumerate(norm_values):
        set_cell_text(row.cells[idx + 2], fmt(value))

# Tabel IV.6.1: normalisasi data latih.
train_norms = {}
train_table = doc.tables[12]
for train in training_rows:
    norm_values = rounded(normalize(train["raw"]))
    train_norms[train["name"]] = norm_values
    row = train_table.rows[train["no"]]
    for idx, value in enumerate(norm_values):
        set_cell_text(row.cells[idx + 2], fmt(value))

# Update x pada paragraf tiap data uji.
x_paragraph_index = {
    "Ridho": 612,
    "FADIL RENALDI SUKMA": 616,
    "Arya Bintang Sipayung": 620,
    "Natasha Kristiani Br Barus": 624,
    "July Arta Gigliati Silaban": 628,
}
for test in test_rows:
    values = test_norms[test["name"]]
    if test["name"] in x_paragraph_index:
        set_paragraph_text(
            doc.paragraphs[x_paragraph_index[test["name"]]],
            "x = (" + ", ".join(fmt(value) for value in values) + ")",
        )

# Hitung ranking jarak untuk semua data uji.
all_rankings = {}
for test in test_rows:
    test_norm = rounded(normalize(test["raw"]))
    ranking = []
    for train in training_rows:
        train_norm = rounded(normalize(train["raw"]))
        full_test_norm = normalize(test["raw"])
        full_train_norm = normalize(train["raw"])
        distance = round(
            sqrt(sum((full_test_norm[idx] - full_train_norm[idx]) ** 2 for idx in range(6))),
            4,
        )
        total_square = round(sum((test_norm[idx] - train_norm[idx]) ** 2 for idx in range(6)), 4)
        ranking.append(
            {
                "no": train["no"],
                "name": train["name"],
                "label": train["label"],
                "norm": train_norm,
                "distance": distance,
                "total_square": total_square,
            }
        )
    ranking.sort(key=lambda item: item["distance"])
    all_rankings[test["name"]] = ranking

# Tabel IV.6.2-IV.6.6: detail perhitungan jarak.
distance_table_indexes = {
    "Ridho": 13,
    "FADIL RENALDI SUKMA": 14,
    "Arya Bintang Sipayung": 15,
    "Natasha Kristiani Br Barus": 16,
    "July Arta Gigliati Silaban": 17,
}
for test in test_rows:
    table = doc.tables[distance_table_indexes[test["name"]]]
    test_norm = test_norms[test["name"]]
    for row_index, item in enumerate(all_rankings[test["name"]], start=1):
        row = table.rows[row_index]
        set_cell_text(row.cells[0], item["no"])
        set_cell_text(row.cells[1], item["name"])
        set_cell_text(row.cells[2], make_formula(test_norm, item["norm"]))
        set_cell_text(row.cells[3], fmt(item["distance"]))

# Tabel IV.7-IV.7.4: lima tetangga terdekat.
top_table_indexes = {
    "Ridho": 18,
    "FADIL RENALDI SUKMA": 19,
    "Arya Bintang Sipayung": 20,
    "Natasha Kristiani Br Barus": 21,
    "July Arta Gigliati Silaban": 22,
}
vote_rows = []
for test in test_rows:
    top_five = all_rankings[test["name"]][:5]
    table = doc.tables[top_table_indexes[test["name"]]]
    for rank, item in enumerate(top_five, start=1):
        row = table.rows[rank]
        set_cell_text(row.cells[0], rank)
        set_cell_text(row.cells[1], item["no"])
        set_cell_text(row.cells[2], item["name"])
        set_cell_text(row.cells[3], fmt(item["distance"]))
        set_cell_text(row.cells[4], item["label"])

    votes = Counter(item["label"] for item in top_five)
    result = votes.most_common(1)[0][0]
    vote_rows.append(
        {
            "no": test["no"],
            "name": test["name"],
            "k1": votes.get("Kelompok Mapel 1", 0),
            "k2": votes.get("Kelompok Mapel 2", 0),
            "result": result,
        }
    )

# Tabel voting.
vote_table = doc.tables[23]
for vote in vote_rows:
    row = vote_table.rows[vote["no"]]
    set_cell_text(row.cells[0], vote["no"])
    set_cell_text(row.cells[1], vote["name"])
    set_cell_text(row.cells[2], vote["k1"])
    set_cell_text(row.cells[3], vote["k2"])
    set_cell_text(row.cells[4], vote["result"])

# Narasi voting.
for idx, vote in enumerate(vote_rows, start=654):
    set_paragraph_text(
        doc.paragraphs[idx],
        f"{vote['no']}. Data uji {vote['name']} memiliki {vote['k1']} suara Kelompok Mapel 1 dan {vote['k2']} suara Kelompok Mapel 2, sehingga hasil rekomendasi sistem adalah {vote['result']}.",
    )

total_k1 = sum(1 for vote in vote_rows if vote["result"] == "Kelompok Mapel 1")
total_k2 = sum(1 for vote in vote_rows if vote["result"] == "Kelompok Mapel 2")
set_paragraph_text(
    doc.paragraphs[659],
    f"Berdasarkan hasil proses voting terhadap kelima data uji, {total_k1} data memperoleh rekomendasi Kelompok Mapel 1 dan {total_k2} data memperoleh rekomendasi Kelompok Mapel 2. Hasil ini menunjukkan bahwa rekomendasi ditentukan melalui kombinasi nilai akademik, Minat Mapel, Bakat Kemampuan, dan kedekatan terhadap data latih.",
)

doc.save(OUTPUT)
print(OUTPUT)
